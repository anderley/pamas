import os
import sys
from pathlib import Path

import django
import mammoth
from docx import Document

BASE_DIR = Path(__file__).resolve().parent.parent
sys.path.append(str(BASE_DIR))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "core.settings")
django.setup()

from quiz.models import Textos


def save_doc(
    tipo: str,
    secao: str,
    nivel: str,
    competencia: str,
    doc: str
):
    print(doc)
    html = None

    if tipo == 'Livro' and competencia:
        style_map = '''
            p[style-name='Heading 1'] => h3:fresh
        '''
        with open(doc, 'rb') as file:
            result = mammoth.convert_to_html(file, style_map=style_map)
            html = result.value
    else:
        competencia = doc.split('/')[-1].split('.')[0]
        document = Document(doc)
        texto = document.tables[0].rows[3].cells[0].text

        tmp_file = '/tmp/texto.docx'

        new_doc = Document()
        new_doc.add_paragraph(texto)
        new_doc.save(tmp_file)

        style_map = '''
            p[style-name='Heading 1'] => h3:fresh
        '''
        with open(tmp_file, 'rb') as file:
            result = mammoth.convert_to_html(file, style_map=style_map)
            html = result.value

    Textos(
        texto=html,
        tipo=tipo,
        secao=secao,
        nivel=nivel,
        competencia=competencia
    ).save()


def main():
    main_dir = 'scripts/Textos'

    for root, d_names, f_names in os.walk(main_dir):
        dirs_names = root.split('/')
        index_tipo = dirs_names.index('Textos') + 1
        tipo = dirs_names[index_tipo] if len(dirs_names) > index_tipo else None # noqa
        index_secao = dirs_names.index('Textos') + 2
        secao = dirs_names[index_secao] if len(dirs_names) > index_secao else None # noqa
        index_nivel = dirs_names.index('Textos') + 3
        nivel = dirs_names[index_nivel] if len(dirs_names) > index_nivel else None # noqa
        index_competencia = dirs_names.index('Textos') + 4
        competencia = dirs_names[index_competencia] if len(dirs_names) > index_competencia else None # noqa

        for f_name in f_names:
            if tipo and secao and nivel and f_names and '.docx' in f_name:
                save_doc(tipo, secao, nivel, competencia, f'{root}/{f_name}')


if __name__ == '__main__':
    main()
